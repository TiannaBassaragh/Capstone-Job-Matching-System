#!/usr/bin/env python3
"""
evaluate_dataset.py — Rank correlation evaluation against annotated vacancy-resume dataset.

Scores 5 vacancies and 30 CVs through the pipeline, then measures how well
the system's rankings agree with two independent human annotators.

Metric: Spearman's ρ per CV, then mean/median/std across all 30 CVs.
Reports results for both match_score and recommendation_score.

Usage:
  cd backend
  python tests/evaluation/evaluate_dataset.py

LLM scores are cached to tests/evaluation/dataset_scores_cache.json so the
script is resumable if interrupted.
"""

import csv
import json
import sys
import time
from pathlib import Path

import numpy as np
from docx import Document
from scipy.stats import spearmanr

# ── path setup ────────────────────────────────────────────────────────────────
BACKEND_DIR  = Path(__file__).resolve().parent.parent.parent
DATASET_DIR  = BACKEND_DIR / "vacancy-resume-matching-dataset-main"
CACHE_FILE   = Path(__file__).parent / "dataset_scores_cache.json"
sys.path.insert(0, str(BACKEND_DIR))

from dotenv import load_dotenv
load_dotenv(BACKEND_DIR / ".env")

from app.database import SessionLocal
from app.models.models import Competency
from app.services.llm_scorer import score_document, extract_tech_keywords
from app.services.scorer import compute_fit_score

# ── human annotations ─────────────────────────────────────────────────────────
# Each row = one CV (CVs 1–30).
# Each cell = rank assigned to that vacancy (1 = best match, 5 = worst).
# Array index: RANKINGS[cv_idx][job_idx], job_idx 0–4 → vacancies 1–5.

ANNOTATOR_1 = [
    [2,1,4,3,5],[1,2,3,4,5],[1,2,3,4,5],[3,1,2,4,5],[1,5,4,2,3],  # CVs 1–5
    [3,2,1,4,5],[3,2,1,5,4],[2,4,3,1,5],[1,5,2,1,4],[3,2,1,4,5],  # CVs 6–10
    [1,2,3,4,5],[1,2,3,4,5],[1,3,2,4,5],[1,2,3,4,5],[3,1,2,4,5],  # CVs 11–15
    [3,1,2,4,5],[3,1,2,4,5],[1,2,5,3,4],[3,2,1,4,5],[3,2,1,4,5],  # CVs 16–20
    [2,3,1,4,5],[1,2,3,5,4],[2,1,3,5,4],[1,2,3,5,4],[1,2,3,4,5],  # CVs 21–25
    [2,1,3,4,5],[2,3,4,5,1],[2,4,3,2,5],[5,1,2,4,3],[2,1,4,3,5],  # CVs 26–30
]

ANNOTATOR_2 = [
    [4,3,1,5,2],[2,4,3,1,5],[5,4,2,3,1],[1,3,2,4,5],[5,1,2,4,3],  # CVs 1–5
    [1,3,2,4,5],[4,2,3,1,5],[2,4,3,1,5],[3,4,2,1,5],[4,1,2,5,3],  # CVs 6–10
    [2,4,3,5,1],[4,3,2,1,5],[4,2,3,1,5],[3,4,2,1,5],[2,4,3,1,5],  # CVs 11–15
    [3,2,4,1,5],[4,2,3,1,5],[4,2,5,3,1],[4,2,3,1,5],[1,5,2,4,3],  # CVs 16–20
    [1,3,4,5,2],[4,1,3,2,5],[1,3,4,2,5],[1,4,3,5,2],[1,4,2,5,3],  # CVs 21–25
    [1,5,2,4,3],[4,3,1,2,5],[1,4,2,3,5],[5,1,2,4,3],[1,2,3,4,5],  # CVs 26–30
]


# ── mock objects (avoids DB writes for the evaluation) ────────────────────────

class _Comp:
    def __init__(self, cid, name, eid):
        self.competency_id   = cid
        self.competency_name = name
        self.onet_element_id = eid


class _JobReq:
    def __init__(self, cid, required_level, importance, requirement_type, comp):
        self.competency_id    = cid
        self.required_level   = required_level
        self.importance       = importance
        self.requirement_type = requirement_type
        self.competency       = comp


# ── helpers ───────────────────────────────────────────────────────────────────

def read_docx(path: Path) -> str:
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    # also extract text from tables (some CVs store content in table cells)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    parts.append(t)
    return "\n".join(parts)


def load_vacancies() -> list[dict]:
    """Return vacancies in CSV row order (= vacancies 1–5 in annotation)."""
    rows = []
    with open(DATASET_DIR / "5_vacancies.csv", newline="", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            rows.append({"title": row["job_title"], "description": row["job_description"]})
    return rows


def build_job_reqs(raw: dict, comp_map: dict, name_map: dict) -> list[_JobReq]:
    reqs = []
    for eid, data in raw.items():
        cid = comp_map.get(eid)
        if cid is None:
            continue
        reqs.append(_JobReq(
            cid               = cid,
            required_level    = data.get("required_level"),
            importance        = data.get("importance"),
            requirement_type  = data.get("requirement_type", "preferred"),
            comp              = _Comp(cid, name_map.get(eid, eid), eid),
        ))
    return reqs


def build_candidate_levels(raw: dict, comp_map: dict) -> dict[int, float | None]:
    return {
        comp_map[eid]: data.get("level")
        for eid, data in raw.items()
        if eid in comp_map
    }


def scores_to_ranks(scores: list[float | None]) -> list[int]:
    """
    [score1..score5] → [rank1..rank5]  (rank 1 = highest score = best match).
    None scores are treated as -inf (ranked last).
    """
    vals  = [s if s is not None else float("-inf") for s in scores]
    order = sorted(range(len(vals)), key=lambda i: vals[i], reverse=True)
    ranks = [0] * len(vals)
    for rank, idx in enumerate(order, start=1):
        ranks[idx] = rank
    return ranks


def print_correlation_table(
    label: str,
    scores_matrix: list[list[float | None]],
    score_key: str = "match_score",
) -> tuple[list[float], list[float], list[float]]:
    """Print per-CV rank correlation table and return (rho_a1, rho_a2, rho_avg)."""
    rho_a1:  list[float] = []
    rho_a2:  list[float] = []
    rho_avg: list[float] = []

    print(f"\n=== Spearman ρ — {label} ===\n")
    print(f"{'CV':>3}  {'System ranks':>15}  {'A1 ranks':>15}  {'A2 ranks':>15}  {'Avg ranks':>17}  {'ρ(A1)':>7}  {'ρ(A2)':>7}  {'ρ(avg)':>7}")
    print("─" * 105)

    for i in range(30):
        sys_ranks = scores_to_ranks(scores_matrix[i])
        a1  = ANNOTATOR_1[i]
        a2  = ANNOTATOR_2[i]
        avg = [round((a1[j] + a2[j]) / 2, 1) for j in range(5)]

        r1,  _ = spearmanr(sys_ranks, a1)
        r2,  _ = spearmanr(sys_ranks, a2)
        ra,  _ = spearmanr(sys_ranks, avg)

        rho_a1.append(r1)
        rho_a2.append(r2)
        rho_avg.append(ra)

        print(
            f"{i+1:>3}  {str(sys_ranks):>15}  {str(a1):>15}  {str(a2):>15}  {str(avg):>17}"
            f"  {r1:>7.3f}  {r2:>7.3f}  {ra:>7.3f}"
        )

    return rho_a1, rho_a2, rho_avg


def print_summary(
    rho_a1: list[float],
    rho_a2: list[float],
    rho_avg: list[float],
    label: str,
) -> None:
    iaa_rhos = [spearmanr(ANNOTATOR_1[i], ANNOTATOR_2[i])[0] for i in range(30)]

    print("\n" + "═" * 75)
    print(f"SUMMARY — {label}")
    print("═" * 75)

    for lbl, rhos in [
        ("vs Annotator 1",   rho_a1),
        ("vs Annotator 2",   rho_a2),
        ("vs Avg Annotators", rho_avg),
    ]:
        arr = np.array(rhos, dtype=float)
        n_valid = int(np.sum(~np.isnan(arr)))
        nan_note = f"  ({30 - n_valid} NaN excluded)" if n_valid < 30 else ""
        print(
            f"  {lbl:22s}  mean ρ = {np.nanmean(arr):+.3f}"
            f"   median = {np.nanmedian(arr):+.3f}"
            f"   std = {np.nanstd(arr):.3f}{nan_note}"
        )

    iaa = np.array(iaa_rhos)
    print()
    print(f"  {'Inter-annotator':22s}  mean ρ = {iaa.mean():+.3f}   median = {np.median(iaa):+.3f}   std = {iaa.std():.3f}")
    print(f"  (Inter-annotator ρ is the human ceiling — our system targets this)")
    print("═" * 75)


# ── main ──────────────────────────────────────────────────────────────────────

def main() -> None:
    db = SessionLocal()
    try:
        all_comps = db.query(Competency).all()
        comp_map  = {c.onet_element_id: c.competency_id   for c in all_comps}
        name_map  = {c.onet_element_id: c.competency_name for c in all_comps}
    finally:
        db.close()

    # load / initialise cache
    cache: dict = {}
    if CACHE_FILE.exists():
        cache = json.loads(CACHE_FILE.read_text())
        print(f"Cache loaded — {len(cache)} entries already scored.")

    def save():
        CACHE_FILE.write_text(json.dumps(cache, indent=2))

    # ── score vacancies ───────────────────────────────────────────────────────
    vacancies = load_vacancies()
    print(f"\n=== Scoring {len(vacancies)} vacancies ===")
    job_reqs_list: list[list[_JobReq]] = []
    job_tech_list:  list[list[str]]    = []

    for i, vac in enumerate(vacancies, start=1):
        key      = f"job_{i}"
        tech_key = f"tech_job_{i}"

        if key not in cache:
            db = SessionLocal()
            try:
                print(f"  [{i}/5] {vac['title']} — scoring...", flush=True)
                raw = score_document(vac["description"], db, mode="job")
                cache[key] = raw
                save()
                time.sleep(0.5)
            finally:
                db.close()
        else:
            print(f"  [{i}/5] {vac['title']} — cached")

        if tech_key not in cache:
            db = SessionLocal()
            try:
                print(f"         → extracting tech keywords...", flush=True)
                tech = extract_tech_keywords(vac["description"], db)
                cache[tech_key] = tech
                save()
                time.sleep(0.3)
            finally:
                db.close()

        reqs = build_job_reqs(cache[key], comp_map, name_map)
        job_reqs_list.append(reqs)
        job_tech_list.append(cache.get(tech_key, []))
        print(f"         → {len(reqs)} competency requirements, {len(cache.get(tech_key, []))} tech keywords")

    # ── score CVs 1–30 ────────────────────────────────────────────────────────
    print(f"\n=== Scoring CVs 1–30 ===")
    candidate_levels_list: list[dict] = []
    candidate_tech_list:   list[list[str]] = []

    for cv_num in range(1, 31):
        key      = f"cv_{cv_num}"
        tech_key = f"tech_cv_{cv_num}"
        cv_path  = DATASET_DIR / "CV" / f"{cv_num}.docx"

        if key not in cache:
            db = SessionLocal()
            try:
                print(f"  [{cv_num}/30] CV {cv_num} — scoring...", flush=True)
                text = read_docx(cv_path)
                if not text.strip():
                    print(f"           ⚠ CV {cv_num} has no extractable text — skipping")
                    cache[key] = {}
                    save()
                else:
                    raw = score_document(text, db, mode="resume")
                    cache[key] = raw
                    save()
                    time.sleep(0.5)
            finally:
                db.close()
        else:
            print(f"  [{cv_num}/30] CV {cv_num} — cached")

        if tech_key not in cache:
            db = SessionLocal()
            try:
                print(f"           → extracting tech keywords...", flush=True)
                text = read_docx(cv_path)
                if text.strip():
                    tech = extract_tech_keywords(text, db)
                    cache[tech_key] = tech
                else:
                    cache[tech_key] = []
                save()
                time.sleep(0.3)
            finally:
                db.close()

        levels = build_candidate_levels(cache[key], comp_map)
        candidate_levels_list.append(levels)
        candidate_tech_list.append(cache.get(tech_key, []))
        print(f"           → {len(levels)} competencies, {len(cache.get(tech_key, []))} tech keywords")

    # ── compute match + recommendation scores ─────────────────────────────────
    print("\n=== Computing match scores (5 × 30 pairs) ===")
    match_scores: list[list[float | None]] = []
    rec_scores:   list[list[float | None]] = []

    for cv_idx, cand_levels in enumerate(candidate_levels_list):
        row_match = []
        row_rec   = []
        for j, job_reqs in enumerate(job_reqs_list):
            result = compute_fit_score(
                cand_levels,
                job_reqs,
                candidate_tech=candidate_tech_list[cv_idx],
                job_tech=job_tech_list[j],
            )
            row_match.append(result["match_score"])
            row_rec.append(result["recommendation_score"])
        match_scores.append(row_match)
        rec_scores.append(row_rec)

    # ── rank correlation — match_score ────────────────────────────────────────
    r1_m, r2_m, ra_m = print_correlation_table("match_score (O*NET fit only)", match_scores)
    print_summary(r1_m, r2_m, ra_m, "match_score")

    # ── rank correlation — recommendation_score ───────────────────────────────
    r1_r, r2_r, ra_r = print_correlation_table("recommendation_score (fit + recall + tech)", rec_scores)
    print_summary(r1_r, r2_r, ra_r, "recommendation_score")

    # ── delta summary ─────────────────────────────────────────────────────────
    print("\n" + "═" * 75)
    print("DELTA — recommendation_score vs match_score (positive = improvement)")
    print("═" * 75)
    for lbl, r_m, r_r in [
        ("vs Annotator 1",    r1_m, r1_r),
        ("vs Annotator 2",    r2_m, r2_r),
        ("vs Avg Annotators", ra_m, ra_r),
    ]:
        delta = np.nanmean(np.array(r_r, dtype=float)) - np.nanmean(np.array(r_m, dtype=float))
        print(f"  {lbl:22s}  Δ mean ρ = {delta:+.3f}")
    print("═" * 75)

    # ── per-vacancy score tables ───────────────────────────────────────────────
    for label, matrix in [("match_score", match_scores), ("recommendation_score", rec_scores)]:
        print(f"\n=== System {label} by CV × vacancy ===")
        print(f"{'CV':>3}  " + "  ".join(f"{'Vac'+str(j+1):>7}" for j in range(5)))
        print("─" * 50)
        for i, row in enumerate(matrix):
            vals = "  ".join(
                f"{v:>7.3f}" if v is not None else f"{'None':>7}"
                for v in row
            )
            print(f"{i+1:>3}  {vals}")


if __name__ == "__main__":
    main()
